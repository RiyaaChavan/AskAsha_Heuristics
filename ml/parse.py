import re
import spacy
import pandas as pd
import docx
import PyPDF2
import os
import json
import argparse
from collections import Counter
from spacy.matcher import PhraseMatcher
from typing import List, Dict, Set, Tuple, Any, Optional, Union

class ResumeSkillExtractor:
    def __init__(self, custom_skills: Optional[Set[str]] = None):
        """
        Initialize the skill extractor with custom skills dictionary and NLP model
        
        Args:
            custom_skills: Optional set of predefined skills to match against
        """
        # Load spaCy model - medium size English model with word vectors
        try:
            self.nlp = spacy.load("en_core_web_md")
        except OSError:
            # If model isn't installed, download it
            import subprocess
            print("Downloading required spaCy model...")
            subprocess.call(["python", "-m", "spacy", "download", "en_core_web_md"])
            self.nlp = spacy.load("en_core_web_md")
        
        # Initialize the default skills dictionary by domain
        self.skills_by_domain = {
            "Technical": {
                "Programming": {"Python", "Java", "JavaScript", "C++", "C#", "Ruby", "Go", "PHP", "Swift", 
                               "Kotlin", "Rust", "Scala", "TypeScript", "R", "Shell", "Perl", "MATLAB"},
                "Web Development": {"HTML", "CSS", "React", "Angular", "Vue.js", "Node.js", "Express", 
                                   "Django", "Flask", "Spring Boot", "ASP.NET", "jQuery", "Bootstrap",
                                   "Sass", "LESS", "WordPress", "Drupal", "Shopify", "Magento", "JAMstack",
                                   "REST APIs", "GraphQL", "WebSockets"},
                "DevOps & Infrastructure": {"Git", "Docker", "Kubernetes", "AWS", "Azure", "Google Cloud", 
                                          "Terraform", "Jenkins", "CircleCI", "Travis CI", "GitHub Actions", 
                                          "Ansible", "Chef", "Puppet", "Nginx", "Apache", "Linux", "Unix", 
                                          "CI/CD", "Infrastructure as Code", "Serverless", "Microservices"},
                "Data Science": {"TensorFlow", "PyTorch", "Pandas", "NumPy", "Scikit-learn", "Keras", 
                                "Natural Language Processing", "NLP", "Computer Vision", "Machine Learning", 
                                "Deep Learning", "Reinforcement Learning", "Data Mining", "Statistical Analysis", 
                                "A/B Testing", "Data Visualization", "Big Data"},
                "Database": {"SQL", "MongoDB", "PostgreSQL", "MySQL", "SQLite", "Oracle", "Redis", 
                            "Cassandra", "DynamoDB", "Elasticsearch", "Neo4j", "MariaDB", "Firestore", 
                            "Database Design", "Data Modeling", "ETL", "Data Warehousing"},
                "Business Intelligence": {"Tableau", "Power BI", "Looker", "QlikView", "Sisense", "Domo", 
                                         "SSRS", "SSIS", "SSAS", "Crystal Reports", "Data Studio"},
                "Cybersecurity": {"Network Security", "Penetration Testing", "Ethical Hacking", "OWASP", 
                                 "Security Auditing", "Vulnerability Assessment", "Encryption", "Firewall Configuration",
                                 "Intrusion Detection", "Security Compliance", "SIEM", "Identity Management"}
            },
            "Finance": {
                "Accounting": {"Financial Accounting", "Managerial Accounting", "Auditing", "Tax Accounting", 
                              "Cost Accounting", "GAAP", "IFRS", "Bookkeeping", "Financial Reporting", 
                              "Accounts Payable", "Accounts Receivable", "General Ledger", "Month-end Close"},
                "Financial Software": {"QuickBooks", "SAP FICO", "NetSuite", "Oracle Financials", "Xero", 
                                     "Sage", "Microsoft Dynamics", "FreshBooks", "Wave", "BlackLine"},
                "Investment": {"Investment Banking", "Portfolio Management", "Asset Management", "Wealth Management",
                              "Private Equity", "Venture Capital", "Hedge Fund", "Derivatives", "Fixed Income", 
                              "Equities", "Securities Analysis", "Financial Modeling", "DCF", "LBO"},
                "Risk & Compliance": {"Risk Management", "Compliance", "AML", "KYC", "Regulatory Reporting", 
                                    "Basel III", "Solvency II", "Dodd-Frank", "MiFID II", "GDPR"}
            },
            "HR & Management": {
                "HR Functions": {"Talent Acquisition", "Recruitment", "HR Analytics", "Employee Engagement",
                               "Performance Management", "Onboarding", "Payroll Management", "Compensation and Benefits",
                               "Labor Law Compliance", "HRIS", "Organizational Development", "Conflict Resolution",
                               "Training and Development", "Succession Planning", "Employee Relations",
                               "Workforce Planning", "Leadership Development"},
                "HR Software": {"Workday", "ADP", "Oracle HCM", "BambooHR", "Zenefits", "SuccessFactors", "Cornerstone",
                              "ATS", "Applicant Tracking Systems", "UltiPro", "Dayforce"},
                "Project Management": {"Agile", "Scrum", "Kanban", "Waterfall", "PRINCE2", "PMBOK", "PMP", 
                                     "Project Planning", "Resource Allocation", "Risk Assessment", "Stakeholder Management"},
                "Management Skills": {"Team Leadership", "Strategic Planning", "Decision Making", "Delegation", 
                                    "Coaching", "Mentoring", "Performance Reviews", "Budget Management", 
                                    "Process Improvement", "Change Management", "Business Strategy"}
            },
            "Marketing & Sales": {
                "Digital Marketing": {"SEO", "SEM", "Content Marketing", "Affiliate Marketing",
                                    "Social Media Marketing", "Email Marketing", "Google Ads", "Facebook Ads", 
                                    "Instagram Ads", "TikTok Marketing", "LinkedIn Marketing", "PPC", "CPC", 
                                    "Conversion Rate Optimization", "A/B Testing", "Marketing Automation"},
                "Marketing Tools": {"Google Analytics", "Google Search Console", "SEMrush", "Ahrefs", "Moz", 
                                  "Mailchimp", "HubSpot", "Marketo", "Salesforce Marketing Cloud", "Adobe Analytics",
                                  "Hootsuite", "Buffer", "Sprout Social", "Canva", "Adobe Creative Suite"},
                "Sales": {"B2B Sales", "B2C Sales", "Inside Sales", "Outside Sales", "Sales Strategy", 
                         "Lead Generation", "Lead Nurturing", "Sales Funnel Management", "Cold Calling", 
                         "Account Management", "Customer Acquisition", "Sales Forecasting"},
                "CRM & Tools": {"Salesforce", "HubSpot CRM", "Zoho CRM", "Microsoft Dynamics CRM", "Sugar CRM", 
                              "Pipedrive", "Close", "Copper", "CRM Administration", "Sales Operations"}
            },
            "Design & Creative": {
                "Design Tools": {"Adobe Photoshop", "Adobe Illustrator", "Adobe XD", "Figma", "Sketch", "InVision",
                               "Adobe InDesign", "Adobe Premiere Pro", "Adobe After Effects", "Canva", 
                               "Blender", "Cinema 4D", "Maya", "ZBrush", "Unity", "Unreal Engine"},
                "Design Skills": {"UI Design", "UX Design", "Wireframing", "Prototyping", "Visual Design", 
                                "Interaction Design", "User Research", "Information Architecture", "Usability Testing",
                                "Responsive Design", "Mobile Design", "Web Design", "Graphic Design",
                                "Logo Design", "Brand Identity", "Typography", "Color Theory"},
                "Other Creative": {"Copywriting", "Content Creation", "Video Production", "Animation", 
                                 "Motion Graphics", "3D Modeling", "Illustration", "Photography", "Storyboarding"}
            },
            "Healthcare": {
                "Clinical": {"Patient Care", "Clinical Assessment", "Medical Diagnosis", "Treatment Planning", 
                            "Medication Management", "Vital Signs Monitoring", "Emergency Response", 
                            "Patient Education", "Physical Examination", "Medical Procedures"},
                "Healthcare IT": {"EHR", "EMR", "Epic", "Cerner", "MEDITECH", "Allscripts", "athenahealth", 
                                "eClinicalWorks", "NextGen", "CPOE", "Health Informatics", "HL7", "FHIR", 
                                "Healthcare Interoperability"},
                "Compliance": {"HIPAA", "HITECH", "Joint Commission", "CMS Regulations", "Medical Coding", 
                             "ICD-10", "CPT Coding", "Medical Billing", "Revenue Cycle Management"},
                "Research": {"Clinical Research", "Medical Writing", "Pharmacovigilance", "Biostatistics", 
                           "Clinical Trials", "IRB Protocols", "GCP", "Literature Review", "Systematic Review"}
            },
            "Soft Skills": {
                "Communication": {"Written Communication", "Verbal Communication", "Public Speaking", 
                                "Presentation Skills", "Active Listening", "Negotiation", "Persuasion",
                                "Conflict Resolution", "Interpersonal Skills", "Client Communication"},
                "Analytical": {"Critical Thinking", "Problem Solving", "Analytical Skills", "Research", 
                              "Data Analysis", "Logical Reasoning", "Quantitative Analysis", "Decision Making"},
                "Work Ethic": {"Time Management", "Organization", "Multitasking", "Prioritization", 
                              "Attention to Detail", "Work Ethic", "Initiative", "Self-Motivation", 
                              "Reliability", "Accountability", "Adaptability", "Flexibility"}
            }
        }
        
        # Create a flat list of all skills from the hierarchical structure
        self.all_skills = set()
        for domain in self.skills_by_domain.values():
            for category in domain.values():
                self.all_skills.update(category)
        
        # Add custom skills if provided
        if custom_skills:
            self.all_skills.update(custom_skills)
        
        # Create spaCy phrase matcher
        self.matcher = PhraseMatcher(self.nlp.vocab, attr="LOWER")
        # Add skill patterns to matcher
        self._add_skill_patterns()
        
        # Regular expressions for skill extraction
        self.skill_patterns = [
            r"(?:^|\W)proficient in\s+([\w\s,\-\(\)/&]+)(?:\W|$)",
            r"(?:^|\W)experienced (?:with|in)\s+([\w\s,\-\(\)/&]+)(?:\W|$)",
            r"(?:^|\W)skills?:?\s+([\w\s,\-\(\)/&]+)(?:\W|$)",
            r"(?:^|\W)knowledge of\s+([\w\s,\-\(\)/&]+)(?:\W|$)",
            r"(?:^|\W)familiar with\s+([\w\s,\-\(\)/&]+)(?:\W|$)",
            r"(?:^|\W)expertise in\s+([\w\s,\-\(\)/&]+)(?:\W|$)",
            r"(?:^|\W)specialized in\s+([\w\s,\-\(\)/&]+)(?:\W|$)",
            r"(?:^|\W)competent (?:with|in)\s+([\w\s,\-\(\)/&]+)(?:\W|$)"
        ]
        self.compiled_patterns = [re.compile(pattern, re.IGNORECASE) for pattern in self.skill_patterns]

    def _add_skill_patterns(self):
        """Add all skills to the PhraseMatcher"""
        skill_patterns = [self.nlp.make_doc(skill) for skill in self.all_skills]
        self.matcher.add("SKILLS", None, *skill_patterns)

    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text content from a PDF file
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text as string
        """
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
        return text

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text content from a DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as string
        """
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
        return text

    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text from a file based on its extension
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text as string
        """
        _, file_extension = os.path.splitext(file_path.lower())
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def extract_skills_with_regex(self, text: str) -> Set[str]:
        """
        Extract skills using regex patterns
        
        Args:
            text: Resume text content
            
        Returns:
            Set of extracted skills
        """
        skills = set()
        
        # Apply regex patterns to extract skills
        for pattern in self.compiled_patterns:
            matches = pattern.findall(text)
            for match in matches:
                # Split by common separators (comma, semicolon)
                items = re.split(r',|;', match)
                for item in items:
                    # Clean up the skill text
                    skill = item.strip()
                    if skill and len(skill) > 2:  # Skip very short terms
                        skills.add(skill)
        
        return skills

    def extract_skills_with_nlp(self, text: str) -> Set[str]:
        """
        Extract skills using NLP and phrase matching
        
        Args:
            text: Resume text content
            
        Returns:
            Set of extracted skills
        """
        skills = set()
        doc = self.nlp(text)
        
        # Use phrase matcher to find skills
        matches = self.matcher(doc)
        for match_id, start, end in matches:
            span = doc[start:end]
            skills.add(span.text)
        
        return skills

    def categorize_skills(self, skills: Set[str]) -> Dict[str, Dict[str, List[str]]]:
        """
        Categorize extracted skills into domains and categories
        
        Args:
            skills: Set of extracted skills
            
        Returns:
            Dictionary of categorized skills
        """
        categorized = {}
        
        # Initialize with empty categories
        for domain, categories in self.skills_by_domain.items():
            categorized[domain] = {}
            for category in categories:
                categorized[domain][category] = []
        
        # Match each skill to its domain and category
        for skill in skills:
            skill_lower = skill.lower()
            found = False
            
            # Try exact match first
            for domain, categories in self.skills_by_domain.items():
                for category, category_skills in categories.items():
                    for category_skill in category_skills:
                        if skill_lower == category_skill.lower() or skill.lower() in category_skill.lower() or category_skill.lower() in skill.lower():
                            categorized[domain][category].append(skill)
                            found = True
                            break
                    if found:
                        break
                if found:
                    break
            
            # If no exact match, use NLP to find the closest category
            if not found:
                skill_doc = self.nlp(skill_lower)
                best_similarity = 0.55  # Threshold for similarity
                best_domain = None
                best_category = None
                
                for domain, categories in self.skills_by_domain.items():
                    for category, category_skills in categories.items():
                        category_doc = self.nlp(category.lower())
                        category_similarity = skill_doc.similarity(category_doc)
                        
                        if category_similarity > best_similarity:
                            best_similarity = category_similarity
                            best_domain = domain
                            best_category = category
                
                if best_domain and best_category:
                    categorized[best_domain][best_category].append(skill)
                else:
                    # If still not categorized, put in "Other" category
                    if "Other" not in categorized:
                        categorized["Other"] = {"Uncategorized": []}
                    if "Uncategorized" not in categorized["Other"]:
                        categorized["Other"]["Uncategorized"] = []
                    categorized["Other"]["Uncategorized"].append(skill)
        
        # Clean up empty categories
        domains_to_remove = []
        for domain, categories in categorized.items():
            categories_to_remove = []
            for category, skills_list in categories.items():
                if not skills_list:
                    categories_to_remove.append(category)
            
            for category in categories_to_remove:
                del categories[category]
            
            if not categories:
                domains_to_remove.append(domain)
        
        for domain in domains_to_remove:
            del categorized[domain]
            
        return categorized

    def extract_skills(self, file_path: str) -> Dict[str, Any]:
        """
        Extract skills from a resume file
        
        Args:
            file_path: Path to the resume file (PDF or DOCX)
            
        Returns:
            Dictionary containing extracted skills and their categorization
        """
        # Extract text from file
        text = self.extract_text_from_file(file_path)
        
        if not text:
            return {"error": "Failed to extract text from file", "skills": [], "categorized": {}}
        
        # Extract skills using different methods
        regex_skills = self.extract_skills_with_regex(text)
        nlp_skills = self.extract_skills_with_nlp(text)
        
        # Combine skills from different methods
        all_extracted_skills = regex_skills.union(nlp_skills)
        
        # Categorize skills
        categorized_skills = self.categorize_skills(all_extracted_skills)
        
        # Count frequency of skills
        skill_frequency = Counter()
        # Look for potential skill mentions throughout the text
        for skill in all_extracted_skills:
            # Count occurrences (case insensitive)
            skill_frequency[skill] = len(re.findall(r'\b' + re.escape(skill) + r'\b', text, re.IGNORECASE))
        
        # Sort skills by frequency
        sorted_skills = sorted(all_extracted_skills, key=lambda x: skill_frequency[x], reverse=True)
        
        return {
            "skills": list(sorted_skills),  # Convert set to list for JSON serialization
            "categorized": categorized_skills,
            "skill_frequency": dict(skill_frequency)
        }

    def to_json(self, result: Dict[str, Any], pretty: bool = True) -> str:
        """
        Convert results dictionary to JSON string
        
        Args:
            result: Result dictionary from extract_skills
            pretty: Whether to format the JSON with indentation
            
        Returns:
            JSON string
        """
        indent = 4 if pretty else None
        return json.dumps(result, indent=indent)
    
    def save_json(self, result: Dict[str, Any], output_path: str, pretty: bool = True) -> None:
        """
        Save results to a JSON file
        
        Args:
            result: Result dictionary from extract_skills
            output_path: Path to save the JSON file
            pretty: Whether to format the JSON with indentation
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=4 if pretty else None)
        print(f"Results saved to {output_path}")


def main():
    """Main function to handle command line arguments and run the skill extractor"""
    parser = argparse.ArgumentParser(description="Extract skills from a resume PDF or DOCX file")
    parser.add_argument("input_file", help="Path to the resume file (PDF or DOCX)")
    parser.add_argument("-o", "--output", help="Path to save the output JSON file (optional)")
    parser.add_argument("--raw", action="store_true", help="Print raw JSON output without indentation")
    parser.add_argument("--summary", action="store_true", help="Print a human-readable summary in addition to JSON")
    
    args = parser.parse_args()
    
    # Initialize with known skills from the prompt
    known_skills = {
        # Tech Skills
        "Python", "Java", "JavaScript", "C++", "SQL", "HTML", "CSS", "React", "Node.js",
        "Django", "Flask", "Git", "Docker", "Kubernetes", "AWS", "Azure", "Google Cloud",
        "TensorFlow", "PyTorch", "Pandas", "NumPy", "Scikit-learn", "Natural Language Processing (NLP)",
        "Computer Vision", "Machine Learning", "Deep Learning", "Tableau", "Power BI",
        "Data Warehousing", "MongoDB", "PostgreSQL", "REST APIs", "GraphQL", "CI/CD",
        "DevOps", "Microservices", "Cybersecurity", "Ethical Hacking", "Selenium", "Jenkins",
        "Hadoop", "Spark", "Blockchain", "Firebase", "Redis", "Kafka",
        # Finance Skills
        "Financial Modeling", "Investment Banking", "Risk Management", "Portfolio Management",
        "Financial Analysis", "Budgeting", "Forecasting", "Auditing", "Accounting", "QuickBooks",
        "SAP FICO", "Equity Research", "Derivatives", "Asset Management", "Wealth Management",
        "Tax Planning", "Corporate Finance", "Mergers and Acquisitions (M&A)", "Treasury Management",
        "Fraud Detection", "Credit Analysis", "Financial Reporting", "IFRS", "GAAP",
        # HR Skills
        "Talent Acquisition", "Recruitment", "HR Analytics", "Employee Engagement",
        "Performance Management", "Onboarding", "Payroll Management", "Compensation and Benefits",
        "Labor Law Compliance", "HRIS", "Organizational Development", "Conflict Resolution",
        "Training and Development (T&D)", "Succession Planning", "Employee Relations",
        "Workforce Planning", "Leadership Development", "ATS",
        # Marketing, Sales, Business Development Skills
        "Digital Marketing", "SEO", "SEM", "Content Marketing", "Affiliate Marketing",
        "Social Media Marketing", "CRM Management", "Brand Management", "Email Marketing",
        "Google Ads", "Facebook Ads", "Influencer Marketing", "Copywriting", "Lead Generation",
        "B2B Sales", "B2C Sales", "Negotiation", "Client Relationship Management",
        "Market Research", "Business Analytics", "Growth Hacking",
        # Healthcare and Life Sciences Skills
        "Clinical Research", "Medical Coding", "Patient Care", "EHR Systems",
        "HIPAA Compliance", "Public Health", "Health Informatics", "Medical Writing",
        "Pharmacovigilance", "Bioinformatics", "Biotechnology",
        # Engineering Skills
        "AutoCAD", "SolidWorks", "ANSYS", "MATLAB", "CATIA", "PLC Programming",
        "SCADA Systems", "Thermodynamics", "Structural Analysis", "Geotechnical Engineering",
        "Circuit Design", "Power Systems", "Embedded Systems",
        # Design Skills
        "Adobe Photoshop", "Adobe Illustrator", "Adobe XD", "Figma", "Sketch", "InVision",
        "Wireframing", "Prototyping", "Visual Storytelling", "Typography", "Branding",
        "Motion Graphics",
        # Other Important Skills (Soft Skills and Certifications)
        "Project Management", "Agile Methodology", "Scrum", "Kanban", "PMP Certification",
        "Six Sigma", "Business Analysis", "Strategic Planning", "Communication Skills",
        "Leadership", "Team Management", "Problem Solving", "Critical Thinking", "Time Management"
    }
    
    try:
        # Create the skill extractor
        skill_extractor = ResumeSkillExtractor(known_skills)
        
        print(f"Processing resume: {args.input_file}")
        
        # Extract skills
        results = skill_extractor.extract_skills(args.input_file)
        
        # Generate JSON output
        pretty = not args.raw
        json_output = skill_extractor.to_json(results, pretty=pretty)
        
        # Print JSON to console
        print("\nExtracted Skills (JSON):")
        print(json_output)
        
        # Save to file if output path is provided
        if args.output:
            skill_extractor.save_json(results, args.output, pretty=pretty)
        
        # Print summary if requested
        if args.summary:
            print("\nSkills Summary:")
            for domain, categories in results["categorized"].items():
                print(f"\n## {domain}")
                for category, skills in categories.items():
                    if skills:
                        print(f"### {category}")
                        for skill in skills:
                            freq = results["skill_frequency"].get(skill, 0)
                            print(f"- {skill} (mentioned {freq} times)")
            
            print(f"\nTotal unique skills found: {len(results['skills'])}")
            
    except Exception as e:
        print(f"Error processing resume: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()